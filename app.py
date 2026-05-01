from __future__ import annotations

import base64
import io
import json
import zipfile

import cv2
import numpy as np
import requests as http_requests
from fastapi import FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from PIL import Image

from backend import history_store
from backend.duplicate_detector import find_duplicates, phash_from_array
from backend.ocr_engine import SUPPORTED_LANGUAGES
from backend.pdf_report import generate_pdf
from backend.pipeline import CulturalHeritagePipeline
from utils.visualization import draw_detections, overlay_mask

app = FastAPI(title="Cultural Heritage Preservation API", version="3.0.0")

ALLOWED_ORIGINS = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "https://heritage-kjqjsvsil-niharxds-projects.vercel.app",
    "https://heritage-r3cnm359a-niharxds-projects.vercel.app",
    "https://*.vercel.app",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

pipeline = CulturalHeritagePipeline()


def read_upload_as_rgb(raw: bytes) -> np.ndarray:
    try:
        return np.array(Image.open(io.BytesIO(raw)).convert("RGB"))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Upload must be a valid image.") from exc


def encode_png_base64(image: np.ndarray) -> str:
    rgb = np.clip(image, 0, 255).astype(np.uint8)
    bgr = cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)
    ok, buffer = cv2.imencode(".png", bgr)
    if not ok:
        raise HTTPException(status_code=500, detail="Could not encode output image.")
    return "data:image/png;base64," + base64.b64encode(buffer).decode("ascii")


def make_thumbnail(image: np.ndarray, size: int = 120) -> str:
    h, w = image.shape[:2]
    scale = size / max(h, w)
    thumb = cv2.resize(image, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_AREA)
    return encode_png_base64(thumb)


def run_pipeline(raw: bytes, domain: str, language: str = "eng") -> dict:
    image  = read_upload_as_rgb(raw)
    result = pipeline.run(image, domain=domain, language=language)

    detection = result["detection"]
    detection_overlay = draw_detections(
        overlay_mask(image, detection["mask"], alpha=0.38),
        detection["detections"],
    )
    heatmap_overlay = overlay_mask(image, result["explanation"]["detection"]["heatmap"], alpha=0.48)

    semantic = result["semantic"]
    if "json" in semantic and isinstance(semantic["json"], str):
        try:
            semantic["labels"] = json.loads(semantic["json"])
        except json.JSONDecodeError:
            pass

    images = {
        "original":         encode_png_base64(image),
        "detectionOverlay": encode_png_base64(detection_overlay),
        "enhanced":         encode_png_base64(result["enhancement"]["image"]),
        "heatmap":          encode_png_base64(heatmap_overlay),
    }

    response = {
        "domain":      result["domain"],
        "images":      images,
        "detection":   {"severity": detection["severity"], "detections": detection["detections"]},
        "enhancement": {"improvement": result["enhancement"]["improvement"], "image": images["enhanced"]},
        "semantic":    semantic,
        "risk":        result["risk"],
        "age":         result["age"],
        "suggestions": result["suggestions"],
        "explanation": result["explanation"]["risk"],
    }

    phash = phash_from_array(image)
    record_id = history_store.save_record(response, make_thumbnail(image), phash=phash)
    response["id"] = record_id
    return response


# ── Health ────────────────────────────────────────────────────────────────────

@app.get("/api/health")
def health() -> dict:
    return {"status": "ok"}


# ── Languages ─────────────────────────────────────────────────────────────────

@app.get("/api/languages")
def get_languages() -> list:
    return [{"name": name, "code": code} for name, code in SUPPORTED_LANGUAGES.items()]


# ── Single Analyze ────────────────────────────────────────────────────────────

@app.post("/api/analyze")
async def analyze(
    file: UploadFile = File(...),
    domain: str = Form("auto"),
    language: str = Form("eng"),
) -> dict:
    if domain not in {"auto", "manuscript", "monument"}:
        raise HTTPException(status_code=400, detail="domain must be auto, manuscript, or monument")
    raw = await file.read()
    return run_pipeline(raw, domain, language)


# ── Batch Analyze ─────────────────────────────────────────────────────────────

@app.post("/api/analyze/batch")
async def analyze_batch(
    files: list[UploadFile] = File(...),
    domain: str = Form("auto"),
    language: str = Form("eng"),
) -> dict:
    if domain not in {"auto", "manuscript", "monument"}:
        raise HTTPException(status_code=400, detail="domain must be auto, manuscript, or monument")
    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 images per batch.")

    results, errors = [], []
    for f in files:
        try:
            raw    = await f.read()
            result = run_pipeline(raw, domain, language)
            results.append({
                "filename":  f.filename,
                "id":        result["id"],
                "domain":    result["domain"],
                "severity":  result["detection"]["severity"],
                "risk":      result["risk"],
                "age":       result["age"],
                "thumbnail": result["images"]["original"],
            })
        except Exception as e:
            errors.append({"filename": f.filename, "error": str(e)})

    return {"results": results, "errors": errors, "total": len(results)}


# ── History ───────────────────────────────────────────────────────────────────

@app.get("/api/history")
def get_history(
    search:     str = Query(""),
    risk_level: str = Query(""),
    domain:     str = Query(""),
    date_from:  str = Query(""),
    date_to:    str = Query(""),
) -> list:
    return history_store.get_all(search=search, risk_level=risk_level, domain=domain, date_from=date_from, date_to=date_to)


@app.get("/api/history/trend")
def get_trend() -> list:
    return history_store.get_trend()


@app.get("/api/history/{record_id}")
def get_record(record_id: str) -> dict:
    record = history_store.get_record(record_id)
    if not record:
        raise HTTPException(status_code=404, detail="Record not found.")
    return record


@app.delete("/api/history/{record_id}")
def delete_record(record_id: str) -> dict:
    if not history_store.delete_record(record_id):
        raise HTTPException(status_code=404, detail="Record not found.")
    return {"deleted": record_id}


# ── Notes ─────────────────────────────────────────────────────────────────────

@app.post("/api/notes/{record_id}")
async def save_notes(record_id: str, payload: dict) -> dict:
    notes = payload.get("notes", "")
    if not history_store.update_notes(record_id, notes):
        raise HTTPException(status_code=404, detail="Record not found.")
    return {"saved": True}


# ── Location & Geocoding ──────────────────────────────────────────────────────

@app.get("/api/geocode")
def geocode(place: str) -> dict:
    """Geocode a place name using Nominatim (OpenStreetMap). No API key needed."""
    try:
        resp = http_requests.get(
            "https://nominatim.openstreetmap.org/search",
            params={"q": place, "format": "json", "limit": 5},
            headers={"User-Agent": "HeritageAI/1.0 (heritage-preservation-tool)"},
            timeout=8,
        )
        results = resp.json()
        return {
            "results": [
                {
                    "place_name": r.get("display_name", ""),
                    "lat":        float(r["lat"]),
                    "lng":        float(r["lon"]),
                }
                for r in results
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Geocoding failed: {e}")


@app.post("/api/location/{record_id}")
async def set_location(record_id: str, payload: dict) -> dict:
    location = {
        "place_name": payload.get("place_name", ""),
        "lat":        payload.get("lat"),
        "lng":        payload.get("lng"),
    }
    if not history_store.update_location(record_id, location):
        raise HTTPException(status_code=404, detail="Record not found.")
    return {"saved": True}


@app.get("/api/map/pins")
def get_map_pins() -> list:
    return history_store.get_map_pins()


# ── Artifact Groups ───────────────────────────────────────────────────────────

@app.get("/api/groups")
def get_groups() -> list:
    return history_store.get_groups()


@app.post("/api/groups")
async def create_group(payload: dict) -> dict:
    name = payload.get("name", "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Group name required.")
    return history_store.create_group(name)


@app.delete("/api/groups/{group_id}")
def delete_group(group_id: str) -> dict:
    if not history_store.delete_group(group_id):
        raise HTTPException(status_code=404, detail="Group not found.")
    return {"deleted": group_id}


@app.post("/api/groups/{group_id}/add/{record_id}")
def add_to_group(group_id: str, record_id: str) -> dict:
    if not history_store.add_to_group(group_id, record_id):
        raise HTTPException(status_code=404, detail="Group not found.")
    return {"added": True}


@app.delete("/api/groups/{group_id}/remove/{record_id}")
def remove_from_group(group_id: str, record_id: str) -> dict:
    history_store.remove_from_group(group_id, record_id)
    return {"removed": True}


@app.get("/api/groups/{group_id}/timeline")
def get_timeline(group_id: str) -> list:
    return history_store.get_group_timeline(group_id)


# ── Settings ──────────────────────────────────────────────────────────────────

@app.get("/api/settings")
def get_settings() -> dict:
    return history_store.get_settings()


@app.post("/api/settings")
async def update_settings(payload: dict) -> dict:
    return history_store.update_settings(payload)


# ── PDF Report ────────────────────────────────────────────────────────────────

@app.get("/api/report/{record_id}")
def download_report(record_id: str) -> Response:
    record = history_store.get_record(record_id)
    if not record:
        raise HTTPException(status_code=404, detail="Record not found.")
    pdf_bytes = generate_pdf(record)
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=heritage_{record_id[:8]}.pdf"},
    )


# ── Urgency Queue ───────────────────────────────────────────────────────────────

@app.get("/api/urgency-queue")
def get_urgency_queue() -> list:
    return history_store.get_urgency_queue()


# ── Checklist ──────────────────────────────────────────────────────────────────

@app.get("/api/checklist/{record_id}")
def get_checklist(record_id: str) -> list:
    return history_store.get_checklist(record_id)


@app.post("/api/checklist/{record_id}")
async def update_checklist(record_id: str, payload: dict) -> list:
    tasks = payload.get("tasks", [])
    return history_store.update_checklist(record_id, tasks)


# ── Duplicate Detection ────────────────────────────────────────────────────────

@app.post("/api/check-duplicate")
async def check_duplicate(file: UploadFile = File(...)) -> dict:
    raw   = await file.read()
    image = read_upload_as_rgb(raw)
    phash = phash_from_array(image)
    stored = history_store.get_all_phashes()
    matches = find_duplicates(phash, stored)
    return {"duplicates": matches, "is_duplicate": len(matches) > 0}


# ── ZIP Export ─────────────────────────────────────────────────────────────────

@app.post("/api/export/zip")
async def export_zip(payload: dict) -> Response:
    record_ids = payload.get("record_ids", [])
    if not record_ids:
        raise HTTPException(status_code=400, detail="No record IDs provided.")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for rid in record_ids:
            record = history_store.get_record(rid)
            if not record:
                continue
            short = rid[:8]
            zf.writestr(f"{short}/analysis.json", json.dumps(record, indent=2))
            try:
                pdf_bytes = generate_pdf(record)
                zf.writestr(f"{short}/report.pdf", pdf_bytes)
            except Exception:
                pass
            thumb_b64 = record.get("thumbnail", "")
            if thumb_b64:
                img_data = base64.b64decode(thumb_b64.split(",", 1)[-1])
                zf.writestr(f"{short}/thumbnail.png", img_data)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=heritage_export.zip"},
    )


# ── DOCX Field Report ─────────────────────────────────────────────────────────

@app.get("/api/docx/{record_id}")
def download_docx(record_id: str) -> Response:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx not installed. Run: pip install python-docx")

    record = history_store.get_record(record_id)
    if not record:
        raise HTTPException(status_code=404, detail="Record not found.")

    full = record.get("full", record)
    doc  = Document()
    doc.add_heading("Heritage Conservation Field Report", 0)
    doc.add_paragraph(f"Record ID: {record_id}")
    doc.add_paragraph(f"Date: {record.get('timestamp', '')}")
    doc.add_paragraph(f"Domain: {record.get('domain', '')}")
    doc.add_heading("Risk Assessment", level=1)
    risk = full.get("risk", {})
    doc.add_paragraph(f"Level: {risk.get('level', '')}   Score: {risk.get('score', 0):.1f}")
    for reason in risk.get("reasons", []):
        doc.add_paragraph(f"• {reason}", style="List Bullet")
    doc.add_heading("Damage Detection", level=1)
    det = full.get("detection", {})
    doc.add_paragraph(f"Severity: {det.get('severity', 0):.1f} / 100")
    for d in det.get("detections", []):
        doc.add_paragraph(f"• {d['label']} — severity {d['severity']:.1f}", style="List Bullet")
    doc.add_heading("Restoration Recommendations", level=1)
    for s in full.get("suggestions", {}).get("suggestions", []):
        doc.add_paragraph(f"[{s['urgency']}] {s['technique']}: {s['description']}", style="List Bullet")
    doc.add_heading("Conservator Notes", level=1)
    doc.add_paragraph(record.get("notes", "") or "(none)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=heritage_{record_id[:8]}.docx"},
    )


# ── QR Code ───────────────────────────────────────────────────────────────────

@app.get("/api/qr/{record_id}")
def get_qr(record_id: str) -> Response:
    try:
        import qrcode
    except ImportError:
        raise HTTPException(status_code=501, detail="qrcode not installed. Run: pip install qrcode[pil]")
    url = f"http://127.0.0.1:5173/?record={record_id}"
    img = qrcode.make(url)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return Response(content=buf.read(), media_type="image/png")


# ── Compare ───────────────────────────────────────────────────────────────────

@app.get("/api/compare/{id_a}/{id_b}")
def compare(id_a: str, id_b: str) -> dict:
    a = history_store.get_record(id_a)
    b = history_store.get_record(id_b)
    if not a or not b:
        raise HTTPException(status_code=404, detail="One or both records not found.")

    fa = a.get("full", a)
    fb = b.get("full", b)

    def _sev(r):   return r.get("detection", {}).get("severity", 0)
    def _score(r): return r.get("risk", {}).get("score", 0)

    return {
        "a": {
            "id": id_a, "timestamp": a["timestamp"], "domain": a["domain"],
            "severity": _sev(fa), "risk": fa.get("risk", {}),
            "age": fa.get("age", {}), "images": fa.get("images", {}),
        },
        "b": {
            "id": id_b, "timestamp": b["timestamp"], "domain": b["domain"],
            "severity": _sev(fb), "risk": fb.get("risk", {}),
            "age": fb.get("age", {}), "images": fb.get("images", {}),
        },
        "diff": {
            "severity_change":   round(_sev(fb)   - _sev(fa),   2),
            "risk_score_change": round(_score(fb) - _score(fa), 2),
        },
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="127.0.0.1", port=8000, reload=True)
